# scripts/make_fixtures.py — run with: python scripts/make_fixtures.py
import fitz
from docx import Document as DocxDocument

doc = fitz.open()
page1 = doc.new_page()
page1.insert_text((72, 100), "Chapter 1: Cell Biology", fontsize=18)
page1.insert_text((72, 140), "The mitochondria is the powerhouse of the cell.", fontsize=11)
page2 = doc.new_page()
page2.insert_text((72, 100), "Chapter 2: Genetics", fontsize=18)
page2.insert_text((72, 140), "DNA carries genetic information in most organisms.", fontsize=11)
doc.save("backend/tests/fixtures/sample.pdf")

docx = DocxDocument()
docx.add_heading("Week 1 Notes", level=1)
docx.add_paragraph("Photosynthesis converts light energy into chemical energy.")
docx.save("backend/tests/fixtures/sample.docx")

scan_doc = fitz.open()
tmp_doc = fitz.open()
tmp_page = tmp_doc.new_page()
tmp_page.insert_text((72, 100), "Chapter 3: Scanned Notes", fontsize=18)
tmp_page.insert_text((72, 140), "Osmosis moves water across a semipermeable membrane.", fontsize=11)
pix = tmp_page.get_pixmap(matrix=fitz.Matrix(2, 2))

scan_page = scan_doc.new_page(width=tmp_page.rect.width, height=tmp_page.rect.height)
scan_page.insert_image(scan_page.rect, pixmap=pix)
scan_doc.save("backend/tests/fixtures/scanned.pdf")
scan_doc.close()
tmp_doc.close()

mixed_doc = fitz.open()
mixed_page1 = mixed_doc.new_page()
mixed_page1.insert_text((72, 100), "Chapter 4: Native Page", fontsize=18)
mixed_page1.insert_text((72, 140), "Ribosomes synthesize proteins from amino acids.", fontsize=11)

tmp_doc2 = fitz.open()
tmp_page2 = tmp_doc2.new_page()
tmp_page2.insert_text((72, 100), "Chapter 5: Scanned Page", fontsize=18)
tmp_page2.insert_text((72, 140), "Chlorophyll absorbs light for photosynthesis.", fontsize=11)
pix2 = tmp_page2.get_pixmap(matrix=fitz.Matrix(2, 2))

mixed_page2 = mixed_doc.new_page(width=tmp_page2.rect.width, height=tmp_page2.rect.height)
mixed_page2.insert_image(mixed_page2.rect, pixmap=pix2)
mixed_doc.save("backend/tests/fixtures/mixed.pdf")
mixed_doc.close()
tmp_doc2.close()
